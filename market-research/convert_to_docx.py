#!/usr/bin/env python3
"""Convert markdown research report to professionally formatted Word document."""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_styles(doc):
    """Create professional styles matching consulting report format."""
    styles = doc.styles

    # Heading 1 - Major sections
    h1 = styles['Heading 1']
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)

    # Heading 2 - Subsections
    h2 = styles['Heading 2']
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 76, 153)

    # Heading 3 - Sub-subsections
    h3 = styles['Heading 3']
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 102, 204)

    # Heading 4
    h4 = styles['Heading 4']
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(51, 51, 51)


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx."""
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i])
        i += 1

    if len(table_lines) < 2:
        return None, start_idx

    header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    rows = []
    for line in table_lines[2:]:
        row = [cell.strip() for cell in line.split('|')[1:-1]]
        if row:
            rows.append(row)

    return {'header': header, 'rows': rows}, i


def add_table(doc, table_data):
    """Add a formatted table to the document."""
    if not table_data or not table_data['header']:
        return

    num_cols = len(table_data['header'])
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(table_data['header']):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row_data in table_data['rows']:
        row = table.add_row()
        for i, cell_text in enumerate(row_data[:num_cols]):
            row.cells[i].text = cell_text
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()


def convert_markdown_to_docx(md_path, docx_path):
    """Convert markdown file to professionally formatted Word document."""

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    doc = Document()
    create_styles(doc)

    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped and i == 0:
            i += 1
            continue

        # Code blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue

        # Tables
        if stripped.startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            table_data, new_i = parse_table(lines, i)
            if table_data:
                add_table(doc, table_data)
                i = new_i
                continue

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:]
            doc.add_heading(text, level=1)
            i += 1
            continue

        if stripped.startswith('## '):
            text = stripped[3:]
            doc.add_heading(text, level=2)
            i += 1
            continue

        if stripped.startswith('### '):
            text = stripped[4:]
            doc.add_heading(text, level=3)
            i += 1
            continue

        if stripped.startswith('#### '):
            text = stripped[5:]
            doc.add_heading(text, level=4)
            i += 1
            continue

        # Blockquotes
        if stripped.startswith('>'):
            quote_text = stripped[1:].strip()
            if quote_text.startswith('>'):
                quote_text = quote_text[1:].strip()

            quote_lines = [quote_text]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith('>'):
                next_quote = lines[j].strip()[1:].strip()
                if next_quote.startswith('>'):
                    next_quote = next_quote[1:].strip()
                quote_lines.append(next_quote)
                j += 1

            full_quote = ' '.join(quote_lines)
            p = doc.add_paragraph(full_quote)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.3)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)

            i = j
            continue

        # Horizontal rules
        if stripped in ['---', '***', '___']:
            doc.add_paragraph('_' * 50)
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Numbered lists
        match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if match:
            text = match.group(2)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Regular paragraphs
        if stripped:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            doc.add_paragraph(text)
            i += 1
            continue

        i += 1

    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python convert_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]

    result = convert_markdown_to_docx(md_path, docx_path)
    print(f"Saved: {result}")
